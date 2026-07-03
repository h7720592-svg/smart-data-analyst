"""Convert Markdown report to Word (.docx) document.

Usage:
    python scripts/convert_report.py <input_md> [output_docx]

Example:
    python scripts/convert_report.py docs/report.md docs/report.docx
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)


def set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading_elm.append(shading)


def add_styled_paragraph(doc, text: str, style: str = "Normal",
                         bold: bool = False, font_size: int = 11,
                         alignment=None, space_after: int = 6,
                         space_before: int = 0, color=None):
    """Add a paragraph with custom formatting."""
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    return para


def convert_markdown_to_docx(md_path: Path, docx_path: Path):
    """Convert a Markdown file to a Word document."""
    if not md_path.exists():
        print(f"Error: Input file not found: {md_path}")
        sys.exit(1)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    i = 0
    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []
    table_alignments: list = []

    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                para = doc.add_paragraph()
                run = para.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                para.paragraph_format.left_indent = Inches(0.3)
                para.paragraph_format.space_after = Pt(8)
                para.paragraph_format.space_before = Pt(4)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            add_styled_paragraph(doc, text, bold=True, font_size=22,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_after=12, color=(0x4F, 0x46, 0xE5))
            i += 1
            continue

        if line.startswith("## "):
            text = line[2:].strip()
            add_styled_paragraph(doc, text, bold=True, font_size=18,
                                 space_after=8, space_before=16,
                                 color=(0x4F, 0x46, 0xE5))
            i += 1
            continue

        if line.startswith("### "):
            text = line[3:].strip()
            add_styled_paragraph(doc, text, bold=True, font_size=14,
                                 space_after=6, space_before=12,
                                 color=(0x1E, 0x29, 0x3B))
            i += 1
            continue

        if line.startswith("|") and line.strip().endswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table_rows = [cells]
                if i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip()):
                    sep_cells = [c.strip() for c in lines[i + 1].split("|")[1:-1]]
                    table_alignments = []
                    for sc in sep_cells:
                        if sc.startswith(":") and sc.endswith(":"):
                            table_alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                        elif sc.endswith(":"):
                            table_alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
                        else:
                            table_alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
                    i += 2
                    continue
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                _render_table(doc, table_rows, table_alignments)
                table_rows = []
                table_alignments = []
                in_table = False

        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = re.sub(r'^[\s]*[-*]\s+', '', line).strip()
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(para, text, 11)
            i += 1
            continue

        if re.match(r'^\s*\d+[\.\)]\s+', line.strip()):
            text = re.sub(r'^\s*\d+[\.\)]\s+', '', line.strip())
            para = doc.add_paragraph(style="List Number")
            _add_formatted_run(para, text, 11)
            i += 1
            continue

        _add_formatted_paragraph(doc, line, 11)
        i += 1

    if in_table and table_rows:
        _render_table(doc, table_rows, table_alignments)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Word document saved: {docx_path}")


def _render_table(doc, rows: list[list[str]], alignments: list):
    """Render a markdown table as a Word table."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols, style="Light Grid Accent 1")
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            if r_idx == 0:
                run.bold = True
                run.font.size = Pt(10)
            else:
                run.font.size = Pt(10)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

            if alignments and c_idx < len(alignments):
                para.alignment = alignments[c_idx]

            if r_idx == 0:
                set_cell_shading(cell, "E8EAF6")

    doc.add_paragraph()


def _add_formatted_paragraph(doc, text: str, font_size: int):
    """Add a paragraph with mixed bold/normal formatting."""
    para = doc.add_paragraph()
    _add_formatted_run(para, text, font_size)
    para.paragraph_format.space_after = Pt(6)


def _add_formatted_run(para, text: str, font_size: int):
    """Parse markdown bold (**text**) and inline code (`text`) in a line."""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(font_size - 1)
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        else:
            part = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', part)
            run = para.add_run(part)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(font_size)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/convert_report.py <input.md> [output.docx]")
        print("Example: python scripts/convert_report.py docs/report.md docs/report.docx")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else input_path.with_suffix(".docx")
    convert_markdown_to_docx(input_path, output_path)
