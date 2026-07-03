"""Generate course design report for Smart Data Analyst project."""

import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUTPUT_DIR = Path("docs")
OUTPUT_PATH = OUTPUT_DIR / "课设报告.docx"

# ── Helpers ────────────────────────────────────────────────────────────────


def set_cell_shading(cell, color: str):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading_elm.append(shading)


def add_paragraph(doc, text, font_size=12, bold=False, alignment=None,
                  space_after=6, space_before=0, font_name="宋体",
                  color=None, first_line_indent=None, line_spacing=1.5):
    """Add a paragraph with full formatting control."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        para.paragraph_format.first_line_indent = first_line_indent
    return para


def add_heading_styled(doc, text, level=1):
    """Add a heading with custom Chinese-friendly styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return heading


def add_table_with_style(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "4F46E5")
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.cell(r + 1, c)
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r % 2 == 0:
                set_cell_shading(cell, "F5F7FA")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacer
    return table


def add_screenshot_placeholder(doc, caption):
    """Add a placeholder box for a screenshot."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)

    # Add a bordered box placeholder
    box_para = doc.add_paragraph()
    box_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = box_para.add_run("〔 此处插入截图 〕")
    run.font.size = Pt(14)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.color.rgb = RGBColor(150, 150, 150)

    # Caption
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_after = Pt(16)
    run = cap_para.add_run(caption)
    run.font.size = Pt(10)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.color.rgb = RGBColor(100, 116, 139)
    run.italic = True


def add_bullet(doc, text, level=0, bold_prefix=""):
    """Add a bullet point."""
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = para.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    para.paragraph_format.line_spacing = 1.5
    return para


# ── Main Report Builder ────────────────────────────────────────────────────


def build_report():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ── Default style ──
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5

    # ==================================================================
    # COVER PAGE
    # ==================================================================

    # School / Course header
    add_paragraph(doc, "华南理工大学", font_size=16, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
                  font_name="黑体", color=(0x1E, 0x29, 0x3B))
    add_paragraph(doc, "开源技术课程设计", font_size=14, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=40,
                  font_name="黑体", color=(0x4F, 0x46, 0xE5))

    # Title
    add_paragraph(doc, "智能数据分析助手", font_size=26, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
                  font_name="黑体", color=(0x1E, 0x29, 0x3B))
    add_paragraph(doc, "Smart Data Analyst", font_size=18, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16,
                  font_name="Times New Roman", color=(0x4F, 0x46, 0xE5))

    # Subtitle
    add_paragraph(doc, "—— 基于 LLM 的智能数据分析与可视化平台", font_size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=60,
                  font_name="楷体", color=(0x64, 0x74, 0x8B))

    # Project info
    info_lines = [
        ("项目地址", "https://github.com/h7720592-svg/smart-data-analyst"),
        ("团队成员", "钟文瀚 (zwhTokisakiKurumi)、黄达杨 (h7720592-svg)"),
        ("指导老师", "__________________"),
        ("提交日期", datetime.now().strftime("%Y 年 %m 月 %d 日")),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        run_label = p.add_run(f"{label}：")
        run_label.font.size = Pt(12)
        run_label.font.name = "宋体"
        run_label._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run_label.bold = True
        run_value = p.add_run(value)
        run_value.font.size = Pt(12)
        run_value.font.name = "宋体"
        run_value._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()

    # ==================================================================
    # ABSTRACT
    # ==================================================================
    add_heading_styled(doc, "摘  要", level=1)

    add_paragraph(doc,
        "随着大数据时代的到来，各行各业积累了海量数据，如何高效、便捷地从数据中提取有价值的信息"
        "成为一项重要挑战。传统的数据分析工具（如 Excel、Python 脚本）虽然功能强大，但对用户的编程"
        "能力和数据科学素养要求较高，普通业务人员往往难以快速上手。近年来，大语言模型（LLM）技术的"
        "飞速发展，为人机交互式的智能数据分析提供了全新的可能性。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_paragraph(doc,
        "本项目设计并实现了一款基于大语言模型的智能数据分析助手（Smart Data Analyst），"
        "采用 Streamlit 轻量级 Web 框架构建用户界面，集成多提供商 LLM（DeepSeek、OpenAI、Groq、"
        "Ollama）作为分析引擎，支持 CSV、Excel、JSON 等多种数据格式的自动加载与画像分析。"
        "用户通过自然语言（中文或英文）与系统对话，AI 即可自动理解分析意图、生成 Python 可视化代码，"
        "并在安全的沙箱环境中执行，最终生成交互式 Plotly 图表和可导出的分析报告（HTML / PDF）。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_paragraph(doc,
        "系统采用三层安全防护机制（AST 白名单校验 + 受限内置函数 + 子进程超时隔离）确保"
        "LLM 生成代码的安全执行。项目完全开源（MIT License），代码托管于 GitHub，"
        "支持社区贡献与二次开发。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_paragraph(doc, "关键词：数据分析；大语言模型；Streamlit；可视化；代码沙箱；开源",
                  font_size=12, bold=True, space_before=12, line_spacing=1.5)

    doc.add_page_break()

    # ==================================================================
    # TABLE OF CONTENTS (placeholder)
    # ==================================================================
    add_heading_styled(doc, "目  录", level=1)
    add_paragraph(doc,
        "（请在 Word 中插入自动目录：引用 → 目录 → 自动目录）",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(150, 150, 150), space_before=40, space_after=20)

    doc.add_page_break()

    # ==================================================================
    # CHAPTER 1: 项目背景与意义
    # ==================================================================
    add_heading_styled(doc, "一、项目背景与意义", level=1)

    add_heading_styled(doc, "1.1 项目背景", level=2)
    add_paragraph(doc,
        "在数据驱动的时代，数据分析能力已成为各行各业的核心竞争力。然而，传统的数据分析流程"
        "存在以下痛点：（1）工具门槛高——Excel 透视表、Python Pandas/Matplotlib 等需要用户具备"
        "一定的编程和统计知识；（2）效率低——从数据清洗到可视化需要反复编写和调试代码；"
        "（3）沟通成本高——业务人员有分析需求但无法直接操作数据，需要与技术人员反复沟通。",
        first_line_indent=Cm(0.74), line_spacing=1.5)
    add_paragraph(doc,
        "2023年以来，以 ChatGPT、DeepSeek 为代表的大语言模型（LLM）展现出强大的代码生成和"
        "自然语言理解能力，为「对话式数据分析」开辟了新的技术路径。用户只需用自然语言描述分析需求，"
        "AI 即可自动生成相应的数据处理和可视化代码，大幅降低数据分析的门槛。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_heading_styled(doc, "1.2 项目意义", level=2)
    add_paragraph(doc,
        "本项目的核心价值在于：（1）降低数据分析门槛——让不具备编程背景的用户也能通过自然语言"
        "对话完成专业级的数据分析；（2）提升分析效率——将原本需要数小时的「理解需求→写代码→调试→"
        "出图」流程压缩到分钟级；（3）保障安全性——通过三层代码沙箱确保 AI 生成的代码不会对系统"
        "造成危害；（4）开源共享——项目以 MIT 协议开源，任何人都可以免费使用、修改和贡献代码，"
        "体现开源技术课程「开放、协作、共享」的核心理念。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_heading_styled(doc, "1.3 课程关联", level=2)
    add_paragraph(doc,
        "本项目是「开源技术」课程的课程设计作品。通过本项目的实践，团队成员深入理解了开源软件的"
        "开发流程，包括：Git 版本控制、GitHub 协作工作流（Fork / Branch / PR / Code Review）、"
        "开源许可证选择（MIT License）、项目文档规范（README / CONTRIBUTING）、CI/CD 自动化测试、"
        "以及社区协作的最佳实践。项目本身就是对课程所学知识的综合运用与展示。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    doc.add_page_break()

    # ==================================================================
    # CHAPTER 2: 需求分析
    # ==================================================================
    add_heading_styled(doc, "二、需求分析", level=1)

    add_heading_styled(doc, "2.1 功能需求", level=2)

    add_paragraph(doc, "（一）数据上传与自动分析", font_size=12, bold=True, space_before=8)
    add_bullet(doc, "支持 CSV（自动编码检测）、Excel (.xlsx/.xls)、JSON 三种数据格式上传")
    add_bullet(doc, "自动检测文件编码（UTF-8 / GBK / GB2312 等），解决中文乱码问题")
    add_bullet(doc, "自动生成数据画像：基本统计量、缺失值比例、异常值检测（IQR 方法）、相关性矩阵")
    add_bullet(doc, "支持大文件处理（上限 200MB），超大数据集自动采样")
    add_bullet(doc, "提供 ydata-profiling 深度数据画像报告导出")

    add_paragraph(doc, "（二）自然语言智能对话", font_size=12, bold=True, space_before=12)
    add_bullet(doc, "支持中文和英文自然语言输入")
    add_bullet(doc, "AI 自动理解分析意图，生成 Python 数据处理与可视化代码")
    add_bullet(doc, "支持 14 种图表类型：柱状图、折线图、散点图、饼图、直方图、箱线图、热力图、"
                        "面积图、小提琴图、密度热力图、旭日图、矩形树图、漏斗图等")
    add_bullet(doc, "内置中文分词（jieba）和基础 NLP 能力，支持文本数据词频分析")
    add_bullet(doc, "代码执行失败时自动纠错重试（最多 2 次）")
    add_bullet(doc, "支持多轮对话上下文记忆")

    add_paragraph(doc, "（三）多提供商 LLM 集成", font_size=12, bold=True, space_before=12)
    add_bullet(doc, "DeepSeek（推荐，性价比高，中文支持好）")
    add_bullet(doc, "OpenAI（GPT-4o / GPT-4o-mini）")
    add_bullet(doc, "Groq（Llama 系列，高速推理）")
    add_bullet(doc, "Ollama（本地部署，数据不出本机）")
    add_bullet(doc, "自定义 OpenAI 兼容 API（支持任意第三方服务）")
    add_bullet(doc, "DeepSeek 深度思考模式（Deep Thinking）")

    add_paragraph(doc, "（四）报告导出", font_size=12, bold=True, space_before=12)
    add_bullet(doc, "HTML 格式报告：交互式图表，可在浏览器中查看和操作")
    add_bullet(doc, "PDF 格式报告：适合打印和正式文档场景")
    add_bullet(doc, "自定义导出内容：可选择包含/排除数据概览、列详情、问题、图表、对话记录等模块")
    add_bullet(doc, "一键下载或本地保存")

    add_heading_styled(doc, "2.2 非功能需求", level=2)
    add_bullet(doc, "安全性：LLM 生成的代码必须在沙箱中执行，不能访问文件系统、网络和系统命令")
    add_bullet(doc, "性能：代码执行超时限制 30 秒，大文件自动采样避免内存溢出")
    add_bullet(doc, "兼容性：支持 Python 3.10+，跨平台（Windows / macOS / Linux）")
    add_bullet(doc, "可维护性：模块化架构，清晰的代码分层，单元测试覆盖核心模块")
    add_bullet(doc, "开源性：MIT 协议，完整的项目文档和贡献指南")

    doc.add_page_break()

    # ==================================================================
    # CHAPTER 3: 系统设计
    # ==================================================================
    add_heading_styled(doc, "三、系统设计", level=1)

    add_heading_styled(doc, "3.1 总体架构", level=2)
    add_paragraph(doc,
        "系统采用经典的三层架构：展示层（Streamlit UI）、业务逻辑层（数据处理 + LLM 集成 + "
        "代码执行）、基础设施层（文件存储 + 报告模板）。各层之间通过明确的接口进行通信，"
        "实现了松耦合、高内聚的模块化设计。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_paragraph(doc, "系统总体架构如下：", first_line_indent=Cm(0.74), line_spacing=1.5)

    # Architecture diagram (text-based)
    arch_text = (
        "┌─────────────────────────────────────────────────────────┐\n"
        "│                    展示层 (Streamlit UI)                  │\n"
        "│  ┌──────────┐  ┌──────────┐  ┌──────────┐              │\n"
        "│  │ 数据上传  │  │ 智能对话  │  │ 导出报告  │              │\n"
        "│  │  Page 1   │  │  Page 2   │  │  Page 3   │              │\n"
        "│  └─────┬─────┘  └─────┬─────┘  └─────┬─────┘              │\n"
        "├────────┼──────────────┼──────────────┼────────────────────┤\n"
        "│        │     业务逻辑层 (Python Backend)     │              │\n"
        "│  ┌─────▼──────────────▼──────────────▼───────────┐        │\n"
        "│  │            Session State Manager               │        │\n"
        "│  ├─────────────────────────────────────────────────┤        │\n"
        "│  │  ┌──────────┐  ┌──────────┐  ┌──────────────┐ │        │\n"
        "│  │  │DataLoader│  │ Profiler │  │  LLM Client  │ │        │\n"
        "│  │  │ CSV/EXCEL│  │  EDA/    │  │  Multi-      │ │        │\n"
        "│  │  │  /JSON   │  │ Issues   │  │  Provider    │ │        │\n"
        "│  │  └──────────┘  └──────────┘  └──────┬───────┘ │        │\n"
        "│  │                       ┌──────────────▼───────┐ │        │\n"
        "│  │                       │   Code Executor      │ │        │\n"
        "│  │                       │ AST Validation +     │ │        │\n"
        "│  │                       │ Sandbox + Timeout    │ │        │\n"
        "│  │                       └──────────┬───────────┘ │        │\n"
        "│  │              ┌───────────────────┼─────────────┐│        │\n"
        "│  │     ┌────────▼───────┐  ┌────────▼───────┐     ││        │\n"
        "│  │     │ Chart Renderer │  │ Report Builder │     ││        │\n"
        "│  │     │  (Plotly)      │  │  (HTML/PDF)    │     ││        │\n"
        "│  │     └────────────────┘  └────────────────┘     ││        │\n"
        "│  └─────────────────────────────────────────────────┘│        │\n"
        "├───────────────────────────────────────────────────────┤        │\n"
        "│             基础设施层 (Storage + Templates)            │        │\n"
        "│   assets/   exports/   profile_reports/   .streamlit/ │        │\n"
        "└─────────────────────────────────────────────────────────┘"
    )
    para = doc.add_paragraph()
    run = para.add_run(arch_text)
    run.font.name = "Consolas"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    para.paragraph_format.space_after = Pt(8)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading_styled(doc, "3.2 模块设计", level=2)

    modules = [
        ("数据加载模块 (data_loader.py)",
         "负责读取 CSV、Excel、JSON 格式文件，自动检测编码（chardet + 多编码回退），"
         "支持自定义分隔符，返回标准化的 DataFrame 和元数据。"),
        ("数据画像模块 (profiler.py)",
         "对加载后的 DataFrame 进行全面的描述性统计分析，包括：基本统计量（均值、中位数、"
         "标准差、分位数等）、缺失值检测、异常值检测（IQR 方法）、偏度分析、重复行检测、"
         "相关性矩阵计算等。"),
        ("LLM 客户端模块 (client.py)",
         "统一的 LLM 调用接口，支持 DeepSeek / OpenAI / Groq / Ollama / 自定义 API 五种"
         "提供商，支持流式输出和结构化 JSON 解析，内置连接验证和错误处理。"),
        ("提示词工程模块 (prompts.py)",
         "定义系统提示词模板，包含数据上下文注入、图表类型描述、代码生成规则、安全约束、"
         "Few-shot 示例，以及错误修复提示词。支持动态构建对话上下文。"),
        ("代码执行沙箱 (code_executor.py)",
         "三层安全防护：AST 白名单校验（仅允许安全模块导入）→ 受限内置函数（屏蔽 open/"
         "exec/eval 等危险函数）→ 子进程超时隔离（multiprocessing + 30s timeout）。"),
        ("可视化模块 (chart_registry.py + renderer.py)",
         "注册 14 种支持的图表类型及其元数据，提供统一的 Plotly 图表渲染接口，内置"
         "一致的视觉主题（配色、字体、布局），支持 HTML 和图片格式导出。"),
        ("报告导出模块 (html_report.py + pdf_report.py)",
         "基于 Jinja2 模板引擎生成 HTML 报告，通过 WeasyPrint 转换为 PDF。支持自定义"
         "报告内容模块（数据概览、列详情、问题、图表、对话记录），内置回退模板。"),
        ("工具函数模块 (utils.py)",
         "提供会话状态管理、日志配置、环境变量加载、API Key 多源获取、数字格式化等通用功能。"),
    ]
    for title, desc in modules:
        add_paragraph(doc, title, font_size=12, bold=True, space_before=10, space_after=2)
        add_paragraph(doc, desc, first_line_indent=Cm(0.74), line_spacing=1.5)

    add_heading_styled(doc, "3.3 技术栈", level=2)
    tech_headers = ["层次", "技术选型", "版本", "选型理由"]
    tech_rows = [
        ["UI 框架", "Streamlit", "≥1.32", "纯 Python 构建 Web 应用，API 简洁，适合数据应用"],
        ["数据处理", "Pandas / NumPy", "≥2.2 / ≥1.26", "业界标准数据分析库，功能完备"],
        ["数据画像", "ydata-profiling", "≥4.8", "一行代码生成完整数据报告"],
        ["可视化", "Plotly", "≥5.18", "交互式图表，支持 30+ 图表类型，导出方便"],
        ["LLM SDK", "OpenAI SDK", "≥1.30", "统一的 OpenAI 兼容接口，多提供商支持"],
        ["中文分词", "jieba", "≥0.42", "轻量级中文分词，无需 GPU"],
        ["编码检测", "chardet", "≥5.2", "自动检测文件编码，解决中文乱码"],
        ["报告模板", "Jinja2", "≥3.1", "灵活的 HTML 模板引擎"],
        ["PDF 导出", "WeasyPrint", "≥61", "纯 Python HTML→PDF 转换"],
        ["测试", "pytest", "≥8.0", "Python 主流测试框架"],
        ["代码检查", "ruff", "≥0.3", "快速 Python linter（替代 flake8）"],
    ]
    add_table_with_style(doc, tech_headers, tech_rows,
                         col_widths=[2.5, 3.0, 2.0, 5.0])

    add_heading_styled(doc, "3.4 安全设计", level=2)
    add_paragraph(doc,
        "由于系统需要执行 LLM 生成的 Python 代码，安全性是最关键的非功能需求。"
        "本系统实现了三层纵深防御体系：",
        first_line_indent=Cm(0.74), line_spacing=1.5)
    add_bullet(doc, "第一层（AST 白名单校验）：在代码执行前解析抽象语法树，仅允许导入白名单中的"
                      "安全模块（pandas、numpy、plotly 等），拦截 os/subprocess/socket 等危险模块，"
                      "同时拦截 open()/exec()/eval()/__import__() 等危险函数调用。",
               bold_prefix="")
    add_bullet(doc, "第二层（受限内置函数）：构建受限的 __builtins__ 字典，移除 open、exec、eval、"
                      "compile、input、breakpoint 等危险内置函数，仅保留安全的数学运算、类型转换等。",
               bold_prefix="")
    add_bullet(doc, "第三层（子进程超时隔离）：代码在独立子进程（multiprocessing.spawn）中执行，"
                      "限时 30 秒。超时自动终止 + 强杀（SIGKILL），防止死循环和资源耗尽。"
                      "子进程无法访问父进程的内存空间。",
               bold_prefix="")

    doc.add_page_break()

    # ==================================================================
    # CHAPTER 4: 系统实现
    # ==================================================================
    add_heading_styled(doc, "四、系统实现", level=1)

    add_heading_styled(doc, "4.1 数据上传与自动分析页面", level=2)
    add_paragraph(doc,
        "用户通过 Streamlit 的 file_uploader 组件上传数据文件（支持拖拽或点击选择），"
        "系统自动检测文件格式和编码。对于 CSV 文件，首先使用 chardet 库进行编码检测，"
        "置信度低于 0.7 时自动回退尝试 UTF-8、GBK、GB2312 等常见中文编码；"
        "同时通过第一行内容自动推断分隔符（逗号、制表符、分号等）。",
        first_line_indent=Cm(0.74), line_spacing=1.5)
    add_paragraph(doc,
        "加载成功后，系统自动执行数据画像分析：计算每列的基本统计量（数值列：均值、标准差、"
        "偏度、分位数、零值比例；类别列：高频值、空字符串比例），并通过 IQR 方法检测异常值。"
        "页面以指标卡片、可展开表格和问题列表三种形式展示结果，用户可一键生成 ydata-profiling "
        "完整数据画像报告。侧边栏提供 LLM 提供商配置（API Key、模型选择、深度思考模式等）。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_heading_styled(doc, "4.2 智能对话与图表生成", level=2)
    add_paragraph(doc,
        "对话页面是系统的核心交互界面。用户输入自然语言问题后，系统首先构建包含数据上下文"
        "（列名、类型、样本值、缺失率等）的系统提示词和对话历史，然后调用 LLM API 获取响应。"
        "LLM 被要求以 JSON 格式返回三种类型之一：text（纯文本回答）、chart（图表代码 + 解释）、"
        "chart+text（图表 + 详细分析）。",
        first_line_indent=Cm(0.74), line_spacing=1.5)
    add_paragraph(doc,
        "对于图表类型请求，系统提取 LLM 生成的 Python 代码，送入安全沙箱执行。"
        "执行成功后，Plotly Figure 对象通过统一主题（配色、字体、布局）渲染为交互式图表，"
        "用户可缩放、悬停查看数据点、下载 PNG 图片。如果代码执行失败，系统自动触发错误修复流程："
        "将错误信息反馈给 LLM，请求修正后重新执行（最多 2 次重试）。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_heading_styled(doc, "4.3 报告导出", level=2)
    add_paragraph(doc,
        "用户在对话分析完成后，可在导出报告页面选择导出格式（HTML / PDF）和包含的内容模块，"
        "一键生成完整的分析报告。HTML 报告基于 Jinja2 模板引擎渲染，嵌入 Plotly 交互式图表"
        "（通过 CDN 加载 plotly.js），支持响应式布局和打印样式。PDF 报告通过 WeasyPrint "
        "将 HTML 转换为 A4 格式，自动配置 CJK 中文字体回退链，支持页码和页眉。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_heading_styled(doc, "4.4 LLM 多提供商集成", level=2)
    add_paragraph(doc,
        "系统通过 OpenAI SDK 的兼容接口实现了统一的多提供商 LLM 调用。五种提供商的配置信息"
        "（Base URL、默认模型列表）集中管理在 PROVIDER_CONFIGS 字典中。用户只需在界面选择"
        "提供商并填入 API Key 即可切换，无需修改代码。对于 DeepSeek 提供商，额外支持了"
        "「深度思考模式」（Deep Thinking），启用后 AI 会进行更深入的推理分析。"
        "自定义提供商选项允许接入任意 OpenAI 兼容 API（如本地 vLLM、Ollama 等）。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_heading_styled(doc, "4.5 项目结构与关键代码规模", level=2)

    # File structure table
    file_headers = ["模块", "文件", "行数", "功能"]
    file_rows = [
        ["应用入口", "app.py", "58", "Streamlit 导航与页面注册"],
        ["数据上传页", "pages/01_数据上传.py", "354", "文件上传、画像、问题检测"],
        ["智能对话页", "pages/02_智能对话.py", "401", "AI 对话、图表生成、自动纠错"],
        ["导出报告页", "pages/03_导出报告.py", "206", "HTML/PDF 报告生成与下载"],
        ["数据加载", "src/data_loader.py", "278", "CSV/Excel/JSON 读取、编码检测"],
        ["数据画像", "src/profiler.py", "276", "统计分析、异常检测、相关性"],
        ["LLM 客户端", "src/llm/client.py", "272", "多提供商 LLM 统一接口"],
        ["提示词工程", "src/llm/prompts.py", "288", "系统提示词、Few-shot 示例"],
        ["代码沙箱", "src/llm/code_executor.py", "436", "AST 校验、沙箱执行、超时控制"],
        ["图表注册", "src/viz/chart_registry.py", "144", "14 种图表类型定义"],
        ["图表渲染", "src/viz/renderer.py", "136", "Plotly 主题化渲染"],
        ["HTML 报告", "src/export/html_report.py", "220", "Jinja2 模板报告生成"],
        ["PDF 报告", "src/export/pdf_report.py", "92", "WeasyPrint PDF 转换"],
        ["工具函数", "src/utils.py", "164", "会话状态、配置、环境变量"],
        ["测试", "tests/ (3 个文件)", "~400", "52 个单元测试"],
        ["报告模板", "assets/report_template.html", "268", "Jinja2 HTML 报告模板"],
    ]
    add_table_with_style(doc, file_headers, file_rows,
                         col_widths=[2.5, 4.0, 1.5, 4.5])
    add_paragraph(doc, f"总计约 3,700 行 Python 代码 + 268 行 HTML 模板 + 52 个测试用例",
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=10, color=(100, 116, 139))

    doc.add_page_break()

    # ==================================================================
    # CHAPTER 5: 系统展示
    # ==================================================================
    add_heading_styled(doc, "五、系统展示", level=1)

    add_paragraph(doc,
        "以下展示系统的两个主要功能界面。运行方式：在项目根目录执行 streamlit run app.py，"
        "浏览器访问 http://localhost:8501。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_paragraph(doc, "5.1 数据上传与画像页面", font_size=12, bold=True, space_before=16)
    add_paragraph(doc,
        "该页面展示数据上传功能及自动分析结果。用户上传数据文件后，页面自动显示数据预览表格、"
        "关键指标卡片（行数、列数、缺失率、重复率、内存占用、数据质量评分）、列详情统计表、"
        "相关性热力图，以及自动检测的数据质量问题列表（缺失值、异常值、偏度、常量列等）。"
        "左侧边栏可配置 LLM 提供商和 API Key。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_screenshot_placeholder(doc, "图 1：数据上传与画像页面截图")

    add_paragraph(doc, "5.2 智能对话与图表生成页面", font_size=12, bold=True, space_before=16)
    add_paragraph(doc,
        "该页面展示自然语言对话分析功能。用户输入分析问题后，AI 自动生成 Python 可视化代码，"
        "在安全沙箱中执行后渲染为交互式 Plotly 图表。页面支持多轮对话上下文记忆，"
        "代码出错时自动纠错重试，图表可交互（缩放、悬停查看数据、下载 PNG）。"
        "左侧边栏提供对话设置、图表类型帮助和清空对话功能，以及 8 个示例问题供快速体验。",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_screenshot_placeholder(doc, "图 2：智能对话与图表生成页面截图")

    doc.add_page_break()

    # ==================================================================
    # CHAPTER 6: 项目分工
    # ==================================================================
    add_heading_styled(doc, "六、项目分工", level=1)

    add_paragraph(doc,
        "本项目由两人协作完成，遵循开源协作模式：项目托管于 GitHub，通过 Git 进行版本控制，"
        "以 Issue 跟踪任务、Pull Request 进行代码合并。两人的具体分工如下：",
        first_line_indent=Cm(0.74), line_spacing=1.5)

    add_paragraph(doc, "6.1 黄达杨（h7720592-svg）—— 后端核心与数据处理", font_size=12,
                  bold=True, space_before=16)

    member1_tasks = [
        ("数据处理模块",
         "data_loader.py（CSV/Excel/JSON 多格式加载、chardet 编码自动检测、分隔符推断）、"
         "profiler.py（描述性统计、IQR 异常值检测、偏度分析、相关性矩阵、ydata-profiling 集成）"),
        ("LLM 集成模块",
         "client.py（多提供商 OpenAI 兼容接口、流式输出、结构化 JSON 解析、连接验证）、"
         "code_executor.py（三层安全沙箱：AST 白名单校验 + 受限内置函数 + 子进程超时隔离）"),
        ("可视化核心",
         "chart_registry.py（14 种图表类型注册与元数据管理）、"
         "renderer.py（Plotly 统一主题渲染、HTML/图片导出接口）"),
        ("应用入口与配置",
         "app.py（Streamlit 多页面导航与全局配置）、"
         "requirements.txt、pyproject.toml、.streamlit/config.toml"),
        ("CI/CD 与测试",
         ".github/workflows/ci.yml（自动化 Lint + 多 Python 版本测试）、"
         "tests/（数据加载、画像分析、代码沙箱共 52 个单元测试）"),
        ("仓库管理",
         "Git 仓库初始化、GitHub 仓库创建、开源发布（MIT License）、"
         "项目结构规划、分支管理与版本控制"),
    ]
    for title, desc in member1_tasks:
        add_paragraph(doc, title, font_size=11, bold=True, space_before=6, space_after=1)
        add_paragraph(doc, desc, font_size=11, first_line_indent=Cm(0.74), line_spacing=1.3)

    add_paragraph(doc, "6.2 钟文瀚（zwhTokisakiKurumi）—— 前端交互与文档报告",
                  font_size=12, bold=True, space_before=16)

    member2_tasks = [
        ("Streamlit 页面开发",
         "pages/01_数据上传.py（文件上传 UI、指标卡片、列详情表格、问题展示、数据画像报告下载）、"
         "pages/02_智能对话.py（聊天界面、LLM 响应解析、图表渲染、自动纠错流程、示例问题）、"
         "pages/03_导出报告.py（导出格式选择、内容模块定制、HTML/PDF 报告生成与下载）"),
        ("前端样式与模板",
         "assets/style.css（指标卡片、问题卡片、聊天消息、按钮等全套 CSS 样式）、"
         "assets/report_template.html（Jinja2 HTML 报告模板，含响应式布局和打印样式）"),
        ("报告导出模块",
         "src/export/html_report.py（Jinja2 报告构建、回退模板）、"
         "src/export/pdf_report.py（WeasyPrint PDF 转换、CJK 字体配置、页码页眉）"),
        ("提示词工程",
         "src/llm/prompts.py（系统提示词设计、Few-shot 示例编写、数据上下文注入、"
         "图表类型描述、代码生成规则约束、错误修复提示词）"),
        ("工具函数与会话管理",
         "src/utils.py（Streamlit 会话状态初始化、日志配置、.env 加载、"
         "API Key 多源获取、数字格式化、LLM 配置构建）"),
        ("项目文档",
         "README.md（项目说明、快速开始、技术栈、安全说明）、"
         "CONTRIBUTING.md（贡献指南、开发流程、PR 规范）、课程设计报告撰写"),
        ("代码优化与改进",
         "修复 API Key 状态处理逻辑、扩展 DeepSeek 模型选项（deepseek-chat/deepseek-reasoner）、"
         "启用流式输出支持、移除冗余代码、更新仓库链接与贡献者信息"),
        ("辅助工具",
         "scripts/convert_report.py（Markdown 转 Word 文档工具）、"
         ".env.example 环境变量模板、.gitignore 配置维护"),
    ]
    for title, desc in member2_tasks:
        add_paragraph(doc, title, font_size=11, bold=True, space_before=6, space_after=1)
        add_paragraph(doc, desc, font_size=11, first_line_indent=Cm(0.74), line_spacing=1.3)

    # Summary table
    add_paragraph(doc, "6.3 分工汇总", font_size=12, bold=True, space_before=16)

    div_headers = ["成员", "GitHub 账号", "主要职责", "核心产出"]
    div_rows = [
        ["黄达杨", "h7720592-svg", "后端核心、数据处理、\n安全沙箱、测试、仓库管理",
         "数据处理模块、LLM 集成、\n代码沙箱、可视化核心、\nCI/CD、52 个测试用例"],
        ["钟文瀚", "zwhTokisakiKurumi", "前端页面、报告导出、\n提示词工程、文档、优化",
         "3 个 Streamlit 页面、\n报告导出模块、CSS 样式、\nHTML 模板、全套文档"],
    ]
    add_table_with_style(doc, div_headers, div_rows,
                         col_widths=[2.0, 3.0, 3.5, 4.0])

    add_paragraph(doc,
        "两人通过 GitHub Pull Request 进行代码审查和合并，使用 Git Commit 记录每次变更。"
        "项目的每一次提交均可追溯至具体贡献者。项目地址："
        "https://github.com/h7720592-svg/smart-data-analyst",
        first_line_indent=Cm(0.74), line_spacing=1.5, space_before=8)

    doc.add_page_break()

    # ==================================================================
    # CHAPTER 7: 总结与展望
    # ==================================================================
    add_heading_styled(doc, "七、总结与展望", level=1)

    add_heading_styled(doc, "7.1 项目总结", level=2)
    add_paragraph(doc,
        "本课程设计项目成功实现了一款基于大语言模型的智能数据分析助手，完成了从需求分析、"
        "系统设计、编码实现、测试验证到开源发布的完整软件开发生命周期。项目的主要成果包括：",
        first_line_indent=Cm(0.74), line_spacing=1.5)
    add_bullet(doc, "一个功能完整的 Web 应用，支持 4 种数据格式、5 种 LLM 提供商、14 种图表类型；")
    add_bullet(doc, "三层代码安全沙箱，确保 AI 生成代码的安全执行；")
    add_bullet(doc, "HTML + PDF 双格式报告导出，满足不同场景需求；")
    add_bullet(doc, "52 个单元测试，覆盖核心模块（数据加载、画像分析、代码沙箱）；")
    add_bullet(doc, "CI/CD 自动化流水线（GitHub Actions），多 Python 版本测试；")
    add_bullet(doc, "完整的开源项目文档（README、CONTRIBUTING、LICENSE）；")
    add_bullet(doc, "模块化的代码架构，清晰的职责分离，便于社区贡献和二次开发。")

    add_heading_styled(doc, "7.2 课程收获", level=2)
    add_paragraph(doc,
        "通过本项目，团队成员在以下方面获得了显著的提升：",
        first_line_indent=Cm(0.74), line_spacing=1.5)
    add_bullet(doc, "深入理解了 Git 分布式版本控制的工作流（branch / commit / merge / rebase）")
    add_bullet(doc, "掌握了 GitHub 开源协作模式（Fork / Clone / PR / Code Review / Issue）")
    add_bullet(doc, "实践了开源许可证的选择与应用（MIT License）")
    add_bullet(doc, "熟悉了 CI/CD 自动化测试的配置与使用（GitHub Actions）")
    add_bullet(doc, "提升了 Python 项目工程化能力（模块化设计、类型注解、单元测试、代码规范）")
    add_bullet(doc, "深入学习了 LLM 应用开发（提示词工程、代码生成、安全沙箱）")
    add_bullet(doc, "体验了从 0 到 1 构建并发布一个完整开源项目的全过程")

    add_heading_styled(doc, "7.3 未来展望", level=2)
    add_paragraph(doc,
        "本项目作为一个课程设计作品，仍有诸多可改进和扩展的方向：",
        first_line_indent=Cm(0.74), line_spacing=1.5)
    add_bullet(doc, "多数据源支持：接入数据库（MySQL / PostgreSQL）、云端存储（Google Sheets、"
                      "Airtable）等数据源，实现数据的实时查询与分析；")
    add_bullet(doc, "高级分析能力：集成统计分析（假设检验、回归分析）、时间序列预测、"
                      "机器学习模型训练等更高级的分析功能；")
    add_bullet(doc, "协作功能：支持多用户协作分析、分析结果分享、对话历史持久化存储；")
    add_bullet(doc, "Docker 部署：提供 Dockerfile 和 docker-compose.yml，简化环境配置和部署流程；")
    add_bullet(doc, "移动端适配：优化 Streamlit 主题和布局，提升移动端浏览体验；")
    add_bullet(doc, "多语言支持：完善国际化（i18n）框架，支持英文、日文等多语言界面。")

    doc.add_page_break()

    # ==================================================================
    # REFERENCES
    # ==================================================================
    add_heading_styled(doc, "参考文献", level=1)

    refs = [
        "[1] Streamlit Documentation. https://docs.streamlit.io/",
        "[2] Plotly Express Documentation. https://plotly.com/python/plotly-express/",
        "[3] Pandas Documentation. https://pandas.pydata.org/docs/",
        "[4] ydata-profiling Documentation. https://docs.profiling.ydata.ai/",
        "[5] OpenAI API Documentation. https://platform.openai.com/docs/",
        "[6] DeepSeek API Documentation. https://platform.deepseek.com/api-docs/",
        "[7] WeasyPrint Documentation. https://doc.courtbouillon.org/weasyprint/",
        "[8] Jinja2 Template Documentation. https://jinja.palletsprojects.com/",
        "[9] Python AST Module Documentation. https://docs.python.org/3/library/ast.html",
        "[10] GitHub Actions Documentation. https://docs.github.com/en/actions",
        "[11] MIT License. https://opensource.org/licenses/MIT",
        "[12] Ruff: An extremely fast Python linter. https://docs.astral.sh/ruff/",
        "[13] Brown, T. B., et al. (2020). Language Models are Few-Shot Learners. NeurIPS 2020.",
        "[14] Chen, M., et al. (2021). Evaluating Large Language Models Trained on Code. arXiv:2107.03374.",
    ]
    for ref in refs:
        add_paragraph(doc, ref, font_size=11, line_spacing=1.5, space_after=4)

    # ==================================================================
    # SAVE
    # ==================================================================
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Report saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
